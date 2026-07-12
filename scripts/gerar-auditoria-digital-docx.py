# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text):
    doc.add_paragraph(text)


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, htxt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = htxt
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = str(val)
    doc.add_paragraph()


h("Auditoria — módulo Digital")
p("Data: 12/07/2026")
p("Fonte: código do repositório (análise read-only).")
p(
    "Escopo: o que é, o que funciona, o que foi implementado, o que falta, SQL, APIs, acesso e riscos."
)
p(
    "Limite: status real das tabelas/colunas no Supabase ao vivo não foi consultado nesta auditoria — indica-se o que o código espera."
)

h("Veredito", 2)
p(
    "O módulo Digital é uma central editorial assistida de redes sociais no admin (/admin/digital):"
)
bullets(
    [
        "cadastro de perfis;",
        "fila de posts;",
        "agente de rascunhos (notícias / eventos / programas);",
        "publicação assistida (copiar/colar) e, opcionalmente, Instagram via Meta Graph API.",
    ]
)
p(
    "Não é um scheduler multi-rede nem uma integração Meta completa. "
    '"Agendado" é lembrete editorial — não publica sozinho.'
)

h("1. Escopo do módulo", 2)
add_table(
    ["Capacidade", "Descrição"],
    [
        [
            "Perfis",
            "Instagram, Facebook, YouTube, LinkedIn, TikTok · escopo site ou projeto",
        ],
        ["Fila", "Status: rascunho → aprovado → agendado → publicado → arquivado"],
        ["Agente", "Gera rascunhos a partir de conteúdo do site; sem publicar"],
        ["Publish assistido", "Copiar texto / abrir rede"],
        ["Instagram Graph", "POST media + media_publish (imagem + legenda)"],
        ["Site público", "/api/public/social-links pode ler digital_accounts"],
    ],
)

h("2. Status por capacidade", 2)
add_table(
    ["Capacidade", "Status", "Notas"],
    [
        ["Perfis (CRUD parcial)", "Pronto", "Sem DELETE na API/UI"],
        ["Fila editorial", "Pronto", "Filtro, resumo, vencidos"],
        [
            "Agente de rascunhos",
            "Pronto",
            "UI + CLI · dedupe por source · sem media_url",
        ],
        ["Agendar", "Lembrete", "Só scheduled_at — sem cron"],
        ["Copiar texto / abrir rede", "Pronto", "Fluxo operacional"],
        [
            "Publicar Instagram (Graph)",
            "Código ok",
            "Exige token + IG id + imagem pública + SQL publish",
        ],
        ["Publicar outras redes", "Ausente", "Só orientação de copiar/colar"],
        [
            "Acesso operador (mod_digital)",
            "Código ok",
            "SQL da coluna precisa estar aplicado no banco",
        ],
        [
            "Links sociais no site",
            "Pronto",
            "Fallback estático se tabela ausente",
        ],
    ],
)

h("3. Arquivos principais", 2)
h("UI admin", 3)
add_table(
    ["Path", "Papel"],
    [
        ["app/admin/digital/page.tsx", "Abas Fila / Perfis"],
        ["app/admin/digital/layout.tsx", "RequireAdminModulo modulo=digital"],
        ["app/admin/layout.tsx", "Menu Digital se pode(digital)"],
        ["app/admin/components/RequireAdminModulo.tsx", "Guard de cliente"],
        ["app/admin/acessos/page.tsx", "Checkbox Digital no escopo"],
    ],
)
h("APIs admin", 3)
add_table(
    ["Path", "Métodos"],
    [
        ["app/api/admin/digital/accounts/route.ts", "GET, POST, PATCH"],
        ["app/api/admin/digital/posts/route.ts", "GET, POST, PATCH"],
        ["app/api/admin/digital/generate/route.ts", "POST"],
        ["app/api/admin/digital/publish/route.ts", "POST (Instagram)"],
        ["app/api/admin/acessos/route.ts", "Lê/grava mod_digital"],
    ],
)
h("API pública", 3)
add_table(
    ["Path", "Papel"],
    [
        ["app/api/public/social-links/route.ts", "Contas site ativas ou fallback"],
        ["components/public/PublicSocialLinks.tsx", "Consome a API"],
    ],
)
h("Lib", 3)
add_table(
    ["Path", "Papel"],
    [
        ["lib/digital/types.ts", "Plataformas, status, tipos"],
        ["lib/digital/labels.ts", "Rótulos PT + ajuda de publicação"],
        ["lib/digital/templates.ts", "Templates notícia/evento/projeto"],
        ["lib/digital/draftAgent.ts", "Agente de rascunhos"],
        ["lib/digital/instagramPublish.ts", "Graph API"],
        ["lib/digital/adminGate.ts", "denyIfSemModuloDigital()"],
        ["lib/auth/adminEscopo.ts", "Módulo digital ↔ mod_digital"],
        [
            "lib/auth/adminSession.ts",
            "Carrega escopos (fallback se coluna ausente)",
        ],
        ["lib/auth/processoLabels.ts", "Chip Digital"],
    ],
)
h("Scripts", 3)
add_table(
    ["Path / comando", "Uso"],
    [
        ["scripts/aplicar-digital-redes-fase1.cjs", "Aplica SQL fase 1"],
        [
            "scripts/aplicar-digital-instagram-publish.cjs",
            "Aplica colunas de publish",
        ],
        ["scripts/aplicar-admin-escopos-mod-digital.cjs", "Aplica mod_digital"],
        ["npm run validar:digital", "Asserts estáticos"],
        ["npm run digital:gerar-rascunhos", "CLI do agente"],
        ["npm run validar:admin-acessos", "Inclui checagem de mod_digital"],
    ],
)

h("4. O que foi implementado", 2)
h("Perfis (digital_accounts)", 3)
bullets(
    [
        "Criar / editar (rótulo, URL, handle, projeto_ref) / ativar-desativar.",
        "Plataformas: Instagram, Facebook, YouTube, LinkedIn, TikTok.",
        "Escopo site ou projeto.",
        "Seed SQL com links institucionais IPECC.",
        "Sem DELETE.",
    ]
)
h("Fila (digital_posts + digital_post_targets)", 3)
bullets(
    [
        "Post manual com destinos (perfis ativos).",
        "Listagem + resumo por status + agendamentos vencidos.",
        "Editar texto/hashtags/mídia/destinos em draft | approved | scheduled.",
        "Aprovar, arquivar, Marcar publicado nas redes.",
        "Copiar texto / copiar link; Abrir plataforma.",
    ]
)
h("Agente", 3)
bullets(
    [
        "Rascunhos de notícias publicadas, eventos, programas.",
        "Deduplica por (source_type, source_id).",
        "Vincula a contas scope=site ativas.",
        "UI + CLI.",
    ]
)
h("Agendamento", 3)
bullets(
    [
        "Status scheduled + scheduled_at.",
        "UI alerta vencidos; lembrete editorial — sem envio automático.",
        "Sem cron/job.",
    ]
)
h("Publicação Instagram", 3)
bullets(
    [
        "Botão para posts approved ou scheduled.",
        "Exige media_url http(s) e destino Instagram ativo.",
        "Grava published_via=instagram_api, external_post_id; em falha, publish_error.",
        "Status final ainda usa published_manual (legado; mitigado por published_via).",
    ]
)
h("Não implementado no código", 3)
bullets(
    [
        "Publicação API Facebook / YouTube / LinkedIn / TikTok.",
        "Stories, Reels, carrossel, vídeo.",
        "Upload de mídia (só URL pública).",
        "Exclusão física de posts/contas.",
        "Isolamento de posts por processo_id (tabelas Digital não têm processo).",
        "OAuth Meta por conta / token por perfil.",
    ]
)

h("5. Controle de acesso", 2)
bullets(
    [
        "Mestre: todos os módulos, incluindo digital.",
        "Operador/externo: só com mod_digital = true no escopo.",
        "UI: RequireAdminModulo; sem permissão redireciona para /admin.",
        "Menu: link só se pode(digital).",
        "APIs: denyIfSemModuloDigital().",
        "Acessos (mestre): checkbox Digital → insert mod_digital.",
    ]
)
p(
    "Atenção de modelo: a flag fica amarrada a um processo em admin_escopos, "
    "mas os dados Digital são globais (sem processo_id). "
    "Operador com Digital vê/edita a fila institucional inteira."
)

h("6. SQL", 2)
add_table(
    ["Arquivo", "Papel"],
    [
        [
            "docs/sql/digital-redes-fase1.sql",
            "Tabelas accounts / posts / targets + seed + RLS",
        ],
        [
            "docs/sql/digital-redes-instagram-publish.sql",
            "Colunas external_post_id, publish_error, published_via",
        ],
        [
            "docs/sql/admin-escopos-mod-digital.sql",
            "Coluna mod_digital em admin_escopos",
        ],
        [
            "docs/sql/multi-admin-processos-fase-2-rls.sql",
            "Case digital em admin_tem_modulo_processo",
        ],
    ],
)
h("Ação pendente (operacional)", 3)
p("Aplicar no SQL Editor do Supabase:")
p(
    "ALTER TABLE public.admin_escopos "
    "ADD COLUMN IF NOT EXISTS mod_digital boolean NOT NULL DEFAULT false;"
)
p("Depois marcar Digital no escopo do operador em /admin/acessos.")

h("7. Integrações Meta / Instagram", 2)
add_table(
    ["Variável", "Uso"],
    [
        ["META_GRAPH_ACCESS_TOKEN", "Token Graph (obrigatório p/ publish)"],
        [
            "INSTAGRAM_BUSINESS_ACCOUNT_ID",
            "IG user id profissional (obrigatório)",
        ],
        ["META_GRAPH_API_VERSION", "Opcional · default v21.0"],
    ],
)
p(
    "Fluxo: POST media (image_url + caption) → media_publish. "
    "Um token global no servidor — não há OAuth por perfil em digital_accounts."
)

h("8. Gaps / backlog", 2)
bullets(
    [
        "Aplicar admin-escopos-mod-digital.sql no Supabase (se ainda não aplicado).",
        "Token Meta + IG Business + imagem pública (go-live Instagram).",
        "Publicação automática no horário agendado.",
        "APIs Facebook / LinkedIn / YouTube / TikTok.",
        "Upload / biblioteca de mídia no admin.",
        "Isolar dados Digital por processo (hoje global).",
        "OAuth / token por conta Instagram no banco.",
        "DELETE de posts/contas; agente preenchendo media_url.",
        "Documentar Digital em docs/ADMIN-SITE-VS-PROCESSO.md.",
        "Testes além dos asserts estáticos de validar:digital.",
    ]
)

h("9. Riscos / débitos técnicos", 2)
bullets(
    [
        "Token Meta global — multi-conta no banco não controla a API.",
        "Status published_manual também para publish via API (mitigado por published_via).",
        "Operador com Digital vê fila inteira — permissão por processo, dados sem isolamento.",
        "Service role nas APIs — quem tem o módulo escreve em todas as digital_*.",
        "Agendamento sem executor — risco de achar que agendado = publicado.",
        "Três migrações SQL — fase1 + instagram-publish + mod_digital.",
        "Instagram exige URL pública de imagem — sem upload, publish costuma ficar bloqueado.",
        "Deduplicação do agente — índice único por source.",
    ]
)

h("10. Resumo", 2)
p(
    "O módulo Digital no código é uma central editorial de redes "
    "(perfis + fila + agente + publish assistido Instagram), integrada ao pacote de "
    "Acessos via mod_digital, com schema e scripts prontos; não é um scheduler "
    "multi-rede nem uma integração Meta completa — Instagram Graph só sobe se "
    "tokens, SQL de publish e imagem pública estiverem no ar."
)

h("Histórico", 2)
p("Commit de escopo Digital no pacote Acessos: 13d2e0e.")
p("Markdown correspondente: docs/AUDITORIA-MODULO-DIGITAL.md")

out = Path(__file__).resolve().parent.parent / "docs" / "AUDITORIA-MODULO-DIGITAL.docx"
doc.save(str(out))
print("OK:", out)
print("Size:", out.stat().st_size, "bytes")
